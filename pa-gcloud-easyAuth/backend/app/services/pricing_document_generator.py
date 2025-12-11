"""
Pricing Document generation service for G-Cloud proposals
Generates Pricing Documents from templates
"""

import os
import logging
from pathlib import Path
from typing import Dict, Optional
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


class PricingDocumentGenerator:
    """Generates G-Cloud Pricing Documents from templates"""
    
    def __init__(self, s3_service=None):
        """
        Initialize pricing document generator
        
        Args:
            s3_service: Optional S3Service instance for AWS Lambda deployment
        """
        self.s3_service = s3_service
        self.use_s3 = s3_service is not None
        
        # Check if we're in Azure (has Azure Storage connection string but no S3)
        self.use_azure = not self.use_s3 and bool(os.environ.get("AZURE_STORAGE_CONNECTION_STRING", ""))
        self.azure_blob_service = None
        if self.use_azure:
            try:
                from app.services.azure_blob_service import AzureBlobService
                self.azure_blob_service = AzureBlobService()
            except Exception as e:
                logger.warning(f"Failed to initialize Azure Blob Service: {e}")
                self.use_azure = False
        
        if self.use_s3:
            # Lambda environment: use /tmp for temporary files
            self.templates_dir = Path("/tmp/templates")
            self.output_dir = Path("/tmp/generated_documents")
        else:
            # Docker/local environment
            self.templates_dir = Path("/app/templates")
            self.output_dir = Path("/app/generated_documents")
        
        # Only create directories if they don't exist and are writable
        try:
            self.output_dir.mkdir(parents=True, exist_ok=True)
        except (OSError, PermissionError) as e:
            if self.use_s3:
                raise
            else:
                # Fallback to /tmp if /app is not writable
                self.templates_dir = Path("/tmp/templates")
                self.output_dir = Path("/tmp/generated_documents")
                self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def generate_pricing_document(
        self,
        service_name: str,
        gcloud_version: str = "15",
        lot: str = "3",
        new_proposal_metadata: Dict | None = None
    ) -> Dict[str, str]:
        """
        Generate Pricing Document from template
        
        Args:
            service_name: Service name/title
            gcloud_version: G-Cloud version (default: "15")
            lot: LOT number (default: "3")
            new_proposal_metadata: Metadata for new proposal (service, owner, sponsor, lot, gcloud_version)
        
        Returns:
            Dict with paths to generated Word document
        """
        
        # Load template
        if self.use_s3:
            # AWS Lambda: download template from S3 to /tmp
            template_key = os.environ.get("PRICING_TEMPLATE_S3_KEY", "templates/pricing_template.docx")
            template_path = self.output_dir / "pricing_template.docx"
            self.s3_service.download_template(template_key, template_path)
        else:
            # Docker/local: use local filesystem
            template_env = os.environ.get("PRICING_TEMPLATE_PATH")
            template_path: Path | None = None
            if template_env:
                env_path = Path(template_env)
                if env_path.exists():
                    template_path = env_path
            
            if template_path is None:
                # Check if we're running in Docker (/app exists) or locally
                is_docker = Path("/app").exists()
                
                if is_docker:
                    # Docker environment: use /app paths
                    docs_dir = Path("/app/docs")
                    candidate = docs_dir / "PA GC15 Pricing Doc SERVICE TITLE.docx"
                    if candidate.exists():
                        template_path = candidate
                    else:
                        template_path = self.templates_dir / "pricing_template.docx"
                else:
                    # Local development: use relative paths from backend directory
                    backend_dir = Path(__file__).parent.parent.parent
                    docs_dir = backend_dir / "docs"
                    templates_dir = backend_dir / "templates"
                    
                    # Check docs first, then templates
                    candidate = docs_dir / "PA GC15 Pricing Doc SERVICE TITLE.docx"
                    if candidate.exists():
                        template_path = candidate
                    elif templates_dir.exists():
                        for p in templates_dir.glob("*pricing*.docx"):
                            template_path = p
                            break
                    
                    if template_path is None:
                        template_path = templates_dir / "pricing_template.docx"
            
            if not Path(template_path).exists():
                raise FileNotFoundError(f"Pricing template not found: {template_path}")
        
        # Create a copy to work with
        doc = Document(str(template_path))
        
        # Replace title with service name - map to Cover Page and Title
        self._replace_title(doc, service_name)
        
        # Replace placeholders across all text nodes
        self._replace_text_in_all_wt(doc, {
            'SERVICE TITLE': service_name,
            'SERVICE NAME': service_name,
            '{{SERVICE_NAME}}': service_name,
        })
        
        # Determine output location and filename
        if new_proposal_metadata:
            # Save to SharePoint folder structure
            service_name_clean = new_proposal_metadata.get('service', service_name)
            lot = new_proposal_metadata.get('lot', lot)
            gcloud_version = new_proposal_metadata.get('gcloud_version', gcloud_version)
            
            # Get folder path using SharePoint service abstraction
            try:
                try:
                    from sharepoint_service.sharepoint_service import get_document_path, USE_S3
                except ImportError:
                    from app.sharepoint_service.sharepoint_service import get_document_path, USE_S3
            except ImportError:
                USE_S3 = False
                get_document_path = None
            
            if self.use_s3 or (get_document_path and USE_S3):
                # S3 environment: folder_path is an S3 prefix (string)
                folder_path = f"GCloud {gcloud_version}/PA Services/Cloud Support Services LOT {lot}/{service_name_clean}/"
            elif self.use_azure:
                # Azure environment: folder_path is a blob prefix (string)
                folder_path = f"GCloud {gcloud_version}/PA Services/Cloud Support Services LOT {lot}/{service_name_clean}/"
            else:
                # Local environment: folder_path is a Path object
                try:
                    try:
                        from sharepoint_service.sharepoint_service import MOCK_BASE_PATH
                    except ImportError:
                        from app.sharepoint_service.sharepoint_service import MOCK_BASE_PATH
                except ImportError:
                    from sharepoint_service.mock_sharepoint import MOCK_BASE_PATH
                
                folder_path = MOCK_BASE_PATH / f"GCloud {gcloud_version}" / "PA Services" / f"Cloud Support Services LOT {lot}" / service_name_clean
            
            # Use exact filename format: PA GC15 Pricing Doc [Service Name].docx
            word_filename = f"PA GC{gcloud_version} Pricing Doc {service_name_clean}.docx"
            
            # Determine where to save the document
            if self.use_s3:
                word_path = self.output_dir / word_filename
                filename_base = service_name_clean
                output_dir = self.output_dir
                s3_key = f"{folder_path}{word_filename}"
            elif self.use_azure:
                word_path = self.output_dir / word_filename
                filename_base = service_name_clean
                output_dir = self.output_dir
                s3_key = None
            else:
                # Local environment: save directly to folder_path (Path object)
                if isinstance(folder_path, str):
                    folder_path = Path(folder_path)
                word_path = folder_path / word_filename
                filename_base = service_name_clean
                output_dir = folder_path
                s3_key = None
        else:
            # Create new document - save to generated_documents (no folder metadata)
            import uuid
            doc_id = str(uuid.uuid4())[:8]
            safe_title = "".join(c for c in service_name if c.isalnum() or c in (' ', '-', '_'))[:50]
            filename_base = f"{safe_title}_pricing_{doc_id}"
            word_path = self.output_dir / f"{filename_base}.docx"
            output_dir = self.output_dir
            s3_key = None
        
        # Ensure output directory exists
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # Save Word document
        doc.save(str(word_path))
        
        # Final safeguard: replace placeholders directly in the saved XML parts
        self._replace_in_saved_docx(str(word_path), {
            'SERVICE TITLE': service_name,
            'SERVICE NAME': service_name,
            '{{SERVICE_NAME}}': service_name,
        })
        
        # Upload to Azure Blob Storage if in Azure environment
        word_blob_key = None
        if self.use_azure and self.azure_blob_service and new_proposal_metadata:
            # Construct blob key matching the SharePoint folder structure
            service_name_clean = new_proposal_metadata.get('service', service_name)
            lot = new_proposal_metadata.get('lot', lot)
            gcloud_version = new_proposal_metadata.get('gcloud_version', gcloud_version)
            word_filename = f"PA GC{gcloud_version} Pricing Doc {service_name_clean}.docx"
            blob_key = f"GCloud {gcloud_version}/PA Services/Cloud Support Services LOT {lot}/{service_name_clean}/{word_filename}"
            
            if blob_key:
                try:
                    self.azure_blob_service.upload_file(word_path, blob_key)
                    word_blob_key = blob_key
                    logger.info(f"Uploaded pricing document to Azure Blob Storage: {word_blob_key}")
                except Exception as e:
                    logger.error(f"Failed to upload pricing document to Azure Blob Storage: {e}")
        
        # Upload to S3 if in Lambda environment
        if self.use_s3 and s3_key:
            import boto3
            s3_client = boto3.client('s3')
            bucket_sharepoint = os.environ.get('SHAREPOINT_BUCKET_NAME', '')
            
            if not bucket_sharepoint:
                raise ValueError("Target S3 bucket for Pricing document not configured")
            
            with open(word_path, 'rb') as f:
                s3_client.upload_fileobj(f, bucket_sharepoint, s3_key)
            
            # Generate presigned URL
            word_url = s3_client.generate_presigned_url(
                'get_object',
                Params={'Bucket': bucket_sharepoint, 'Key': s3_key},
                ExpiresIn=3600
            )
            
            return {
                "word_path": word_url,
                "word_s3_key": s3_key,
                "filename": filename_base
            }
        elif self.use_azure:
            # Azure: return blob keys
            return {
                "word_path": str(word_path),
                "word_blob_key": word_blob_key,
                "filename": filename_base
            }
        else:
            # Docker/local: return local paths
            return {
                "word_path": str(word_path),
                "filename": filename_base
            }
    
    def _replace_title(self, doc: Document, service_name: str):
        """Replace the first Heading 1 with service name, or append to existing title"""
        replaced = False
        for paragraph in doc.paragraphs:
            if paragraph.style.name == 'Heading 1':
                # Found the title - append service name to existing title
                # Format: "PA CONSULTING GCLOUD PRICING DOCUMENT - [Service Name]"
                current_text = paragraph.text.strip()
                if service_name not in current_text:
                    paragraph.text = f"{current_text} - {service_name}"
                else:
                    paragraph.text = current_text
                # Preserve formatting
                for run in paragraph.runs:
                    run.font.size = Pt(24)
                    run.font.bold = True
                replaced = True
                break
        
        # Fallback: replace literal occurrences of placeholders
        if not replaced:
            placeholders = ['SERVICE TITLE', 'SERVICE NAME', '{{SERVICE_NAME}}']
            for p in doc.paragraphs:
                for r in p.runs:
                    for placeholder in placeholders:
                        if placeholder in r.text:
                            r.text = r.text.replace(placeholder, service_name)
                            replaced = True
                            break
                    if replaced:
                        break
                if replaced:
                    break
    
    def _replace_text_in_all_wt(self, doc: Document, mapping: Dict[str, str]):
        """Replace text in all w:t nodes within the main document"""
        ns = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        }
        try:
            for xpath in ('.//w:t', './/a:t'):
                for t in doc._element.xpath(xpath, namespaces=ns):
                    if t.text:
                        txt = t.text
                        replaced = txt
                        for old, new in mapping.items():
                            if old in replaced:
                                replaced = replaced.replace(old, new)
                        if replaced != txt:
                            t.text = replaced
        except Exception:
            pass
    
    def _replace_in_saved_docx(self, docx_path: str, mapping: Dict[str, str]):
        """Open the saved .docx and replace placeholders in XML files as a last step"""
        try:
            from zipfile import ZipFile, ZIP_DEFLATED
            import io
            with ZipFile(docx_path, 'r') as zin:
                buf = io.BytesIO()
                with ZipFile(buf, 'w', ZIP_DEFLATED) as zout:
                    for item in zin.infolist():
                        data = zin.read(item.filename)
                        if item.filename.startswith('word/') and item.filename.endswith('.xml'):
                            try:
                                text = data.decode('utf-8')
                                for old, new in mapping.items():
                                    if old in text:
                                        text = text.replace(old, new)
                                data = text.encode('utf-8')
                            except Exception:
                                pass
                        zout.writestr(item, data)
            # Overwrite original file
            with open(docx_path, 'wb') as f:
                f.write(buf.getvalue())
        except Exception:
            pass


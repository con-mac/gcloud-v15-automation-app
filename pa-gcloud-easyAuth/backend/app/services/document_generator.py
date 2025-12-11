"""
Document generation service for G-Cloud proposals
Generates Word and PDF documents from templates
"""

import os
import json
import shutil
import logging
from pathlib import Path
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
import uuid
from bs4 import BeautifulSoup
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


class DocumentGenerator:
    """Generates G-Cloud proposal documents from templates"""
    
    def __init__(self, s3_service=None):
        """
        Initialize document generator
        
        Args:
            s3_service: Optional S3Service instance for AWS Lambda deployment
        """
        self.s3_service = s3_service
        self.use_s3 = s3_service is not None
        
        # Check if we're in Azure (has Azure Storage connection string but no S3)
        self.use_azure = not self.use_s3 and bool(os.environ.get("AZURE_STORAGE_CONNECTION_STRING", ""))
        self.azure_blob_service = None
        if self.use_azure:
            try:
                from app.services.azure_blob_service import AzureBlobService
                self.azure_blob_service = AzureBlobService()
            except Exception as e:
                logger.warning(f"Failed to initialize Azure Blob Service: {e}")
                self.use_azure = False
        
        if self.use_s3:
            # Lambda environment: use /tmp for temporary files
            self.templates_dir = Path("/tmp/templates")
            self.output_dir = Path("/tmp/generated_documents")
        else:
            # Docker/local environment
            self.templates_dir = Path("/app/templates")
            self.output_dir = Path("/app/generated_documents")
        
        # Only create directories if they don't exist and are writable
        # In Lambda, /tmp is always available, but we don't need to create /app
        try:
            self.output_dir.mkdir(parents=True, exist_ok=True)
        except (OSError, PermissionError) as e:
            # If we can't create directories (e.g., Lambda read-only filesystem),
            # fall back to /tmp (which is always writable in Lambda)
            if self.use_s3:
                # Already using /tmp, so this shouldn't happen
                raise
            else:
                # Fallback to /tmp if /app is not writable
                self.templates_dir = Path("/tmp/templates")
                self.output_dir = Path("/tmp/generated_documents")
                self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def generate_service_description(
        self,
        title: str,
        description: str,
        features: List[str],
        benefits: List[str],
        service_definition: List[dict] | None = None,
        update_metadata: Dict | None = None,
        save_as_draft: bool = False,
        new_proposal_metadata: Dict | None = None
    ) -> Dict[str, str]:
        """
        Generate Service Description document from template
        
        Args:
            title: Service name/title
            description: Short service description (500 words max)
            features: List of service features (10 words each, max 10)
            benefits: List of service benefits (10 words each, max 10)
        
        Returns:
            Dict with paths to generated Word and PDF files
        """
        
        # Load template
        if self.use_s3:
            # AWS Lambda: download template from S3 to /tmp
            template_key = os.environ.get("TEMPLATE_S3_KEY", "templates/service_description_template.docx")
            template_path = self.output_dir / "template.docx"
            self.s3_service.download_template(template_key, template_path)
        else:
            # Docker/local: use local filesystem
            template_env = os.environ.get("SERVICE_DESC_TEMPLATE_PATH")
            template_path: Path | None = None
            if template_env:
                env_path = Path(template_env)
                if env_path.exists():
                    template_path = env_path
            
            if template_path is None:
                # Check if we're running in Docker (/app exists) or locally
                is_docker = Path("/app").exists()
                
                if is_docker:
                    # Docker environment: use /app paths
                    docs_dir = Path("/app/docs")
                    candidate = None
                    if docs_dir.exists():
                        for p in docs_dir.glob("*.docx"):
                            candidate = p
                            break
                    template_path = candidate or (self.templates_dir / "service_description_template.docx")
                else:
                    # Local development: use relative paths from backend directory
                    # Get the backend directory (parent of app directory)
                    backend_dir = Path(__file__).parent.parent.parent
                    templates_dir = backend_dir / "templates"
                    docs_dir = backend_dir / "docs"
                    
                    # Check docs first, then templates
                    candidate = None
                    if docs_dir.exists():
                        for p in docs_dir.glob("*.docx"):
                            candidate = p
                            break
                    if candidate is None and templates_dir.exists():
                        for p in templates_dir.glob("*.docx"):
                            candidate = p
                            break
                    
                    if candidate is None:
                        # Fallback to templates directory
                        template_path = templates_dir / "service_description_template.docx"
                    else:
                        template_path = candidate
            
            if not Path(template_path).exists():
                raise FileNotFoundError(f"Template not found: {template_path}")
        
        # Create a copy to work with
        doc = Document(str(template_path))
        
        # Replace title (first Heading 1)
        self._replace_title(doc, title)
        
        # Clean target sections entirely and build a fresh content block after TOC
        self._remove_sections(doc, [
            'Short Service Description',
            'Key Service Features',
            'Key Service Benefits',
            'Service Definition',
        ])
        # Capture and remove 'About PA' block using marker or heading, to re-append at the very end
        about_pa_block = self._extract_block_by_marker(doc, '{{ABOUT_PA_START}}')
        if not about_pa_block:
            about_pa_block = self._extract_heading_block(doc, 'About PA')

        after_toc_para = self._ensure_toc_and_pagebreak(doc)
        # Insert fresh content block starting on a new page
        last_para = self._insert_full_content_block(
            doc=doc,
            after_para=after_toc_para,
            service_title=title,
            description=description,
            features=features,
            benefits=benefits,
            service_definition=service_definition or [],
        )
        # Re-append About PA block as the final page on its own
        if about_pa_block:
            # Page break before About PA - ensure proper spacing
            # Add multiple blank paragraphs and explicit page break for proper separation
            tail_para = last_para or (doc.paragraphs[-1] if doc.paragraphs else None)
            if tail_para is not None:
                # Add a blank paragraph for spacing
                blank_para1 = self._insert_paragraph_after(tail_para, '')
                # Add another blank paragraph
                blank_para2 = self._insert_paragraph_after(blank_para1, '')
                # Add page break to the second blank paragraph (ensures proper page break)
                blank_para2.add_run().add_break(WD_BREAK.PAGE)
                # Add one more blank paragraph after page break to ensure spacing
                blank_para3 = self._insert_paragraph_after(blank_para2, '')
            body = doc._element.body
            for el in about_pa_block:
                body.append(el)

        # Placeholders and ToC handling occur after content is built
        self._enable_update_fields_on_open(doc)
        
        # Update TOC field after content insertion
        # This ensures the TOC is populated with actual headings
        self._update_toc_field(doc)

        # Replace placeholders across all text nodes (including inside shapes/textboxes)
        self._replace_text_in_all_wt(doc, {
            'ENTER SERVICE NAME HERE': title,
            'Enter Service Name Here': title,
            'enter service name here': title,
            'Add Title': title,
            '{{SERVICE_NAME}}': title,
        })
        
        # Determine output location and filename
        if update_metadata or new_proposal_metadata:
            # Save to SharePoint folder (either update or new proposal)
            if update_metadata:
                folder_path = update_metadata.get('folder_path') or ''
                gcloud_version = update_metadata.get('gcloud_version', '14')
                doc_type = update_metadata.get('doc_type', 'SERVICE DESC')
                service_name = update_metadata.get('service_name', title)
                actual_folder_name = service_name  # Default to service_name
                
                # If folder_path is empty and we're in Azure/S3, construct it
                if not folder_path:
                    if self.use_azure:
                        folder_path = f"GCloud {gcloud_version}/PA Services/Cloud Support Services LOT {update_metadata.get('lot', '2')}/{service_name}/"
                    elif self.use_s3:
                        folder_path = f"GCloud {gcloud_version}/PA Services/Cloud Support Services LOT {update_metadata.get('lot', '2')}/{service_name}/"
                
                # Only convert to Path for local filesystem (not Azure/S3)
                if not self.use_s3 and not self.use_azure and folder_path:
                    folder_path = Path(folder_path)
                    # Extract actual folder name from path for filename consistency
                    if isinstance(folder_path, Path) and folder_path.exists():
                        actual_folder_name = folder_path.name
                elif self.use_azure and folder_path:
                    # For Azure, folder_path is a blob prefix string like "GCloud 15/PA Services/Cloud Support Services LOT 2/Service Name/"
                    # Extract service folder name from the blob prefix
                    parts = folder_path.rstrip('/').split('/')
                    if len(parts) > 0:
                        # Last non-empty part is the service folder name
                        for part in reversed(parts):
                            if part:
                                actual_folder_name = part
                                break
            else:  # new_proposal_metadata
                service_name = new_proposal_metadata.get('service', title)
                lot = new_proposal_metadata.get('lot', '2')
                gcloud_version = new_proposal_metadata.get('gcloud_version', '15')
                doc_type = 'SERVICE DESC'  # Always SERVICE DESC for new proposals
                
                # Get folder path using SharePoint service abstraction
                try:
                    try:
                        from sharepoint_service.sharepoint_service import get_document_path, USE_S3
                    except ImportError:
                        from app.sharepoint_service.sharepoint_service import get_document_path, USE_S3
                except ImportError:
                    USE_S3 = False
                    get_document_path = None
                
                # Initialize actual_folder_name - will be used for filename regardless of environment
                actual_folder_name = service_name  # Default to service_name
                
                if self.use_s3 or (get_document_path and USE_S3):
                    # S3 environment: folder_path is an S3 prefix (string)
                    # Construct S3 prefix: GCloud {version}/PA Services/Cloud Support Services LOT {lot}/{service_name}/
                    folder_path = f"GCloud {gcloud_version}/PA Services/Cloud Support Services LOT {lot}/{service_name}/"
                    # For S3, actual_folder_name stays as service_name (already set above)
                elif self.use_azure:
                    # Azure environment: folder_path is a blob prefix (string)
                    # Use service_name directly with spaces - NO normalization
                    folder_path = f"GCloud {gcloud_version}/PA Services/Cloud Support Services LOT {lot}/{service_name}/"
                    # For Azure, actual_folder_name stays as service_name (already set above)
                else:
                    # Local environment: folder_path is a Path object
                    # But if MOCK_BASE_PATH is None (Azure environment), fallback to output_dir
                    try:
                        try:
                            from sharepoint_service.sharepoint_service import MOCK_BASE_PATH
                        except ImportError:
                            from app.sharepoint_service.sharepoint_service import MOCK_BASE_PATH
                    except ImportError:
                        from sharepoint_service.mock_sharepoint import MOCK_BASE_PATH
                    
                    # Check if MOCK_BASE_PATH is None (Azure environment without local mock)
                    if MOCK_BASE_PATH is None:
                        # Fallback to output_dir if MOCK_BASE_PATH is None
                        logger.warning("MOCK_BASE_PATH is None, using output_dir for folder_path")
                        folder_path = self.output_dir
                        actual_folder_name = service_name
                    else:
                        folder_path = MOCK_BASE_PATH / f"GCloud {gcloud_version}" / "PA Services" / f"Cloud Support Services LOT {lot}" / service_name
                        
                        # Verify folder exists, if not try to find it with fuzzy match
                        actual_folder_name = service_name
                        if not folder_path.exists():
                            # Try to find folder with fuzzy match
                            base_path = MOCK_BASE_PATH / f"GCloud {gcloud_version}" / "PA Services"
                            lot_folder = base_path / f"Cloud Support Services LOT {lot}"
                            if lot_folder.exists():
                                from sharepoint_service.mock_sharepoint import fuzzy_match
                                for folder in lot_folder.iterdir():
                                    if folder.is_dir() and fuzzy_match(service_name, folder.name):
                                        folder_path = folder
                                        actual_folder_name = folder.name  # Use actual folder name for filename
                                        break
            
            # Use exact filename format: PA GC15 SERVICE DESC [Folder Name].docx
            # Use actual_folder_name (folder.name) to match what get_document_path expects
            # This ensures saved files can be found when reopening
            if doc_type == 'SERVICE DESC':
                word_filename = f"PA GC{gcloud_version} SERVICE DESC {actual_folder_name}.docx"
            else:
                word_filename = f"PA GC{gcloud_version} Pricing Doc {actual_folder_name}.docx"
            
            # Add _draft suffix if saving as draft
            if save_as_draft:
                word_filename = word_filename.replace('.docx', '_draft.docx')
            else:
                # Remove any existing _draft files when completing (only for local)
                # Ensure folder_path is valid before using it
                if not self.use_s3 and not self.use_azure:
                    # Ensure folder_path is a Path object
                    if folder_path is None:
                        folder_path = self.output_dir
                    elif isinstance(folder_path, str):
                        folder_path = Path(folder_path)
                    elif not isinstance(folder_path, Path):
                        folder_path = self.output_dir
                    
                    if isinstance(folder_path, Path) and folder_path.exists():
                        draft_filename = word_filename.replace('.docx', '_draft.docx')
                        draft_path = folder_path / draft_filename
                        if draft_path.exists():
                            draft_path.unlink()
                        
                        # Remove any existing SERVICE DESC files (to replace them)
                        if 'SERVICE DESC' in word_filename:
                            for existing_file in folder_path.glob(f"PA GC{gcloud_version} SERVICE DESC {service_name}*.docx"):
                                if existing_file.name != word_filename:
                                    existing_file.unlink()
            
            # Determine where to save the document
            if self.use_s3:
                # S3 environment: save to /tmp first, then upload to S3
                # folder_path is an S3 prefix (string), word_filename is the filename
                word_path = self.output_dir / word_filename  # Save to /tmp first
                filename_base = service_name
                output_dir = self.output_dir
                s3_key = f"{folder_path}{word_filename}"  # Full S3 key
            elif self.use_azure:
                # Azure environment: save to /tmp first, then upload to Azure Blob Storage
                # folder_path is a blob prefix (string), word_filename is the filename
                word_path = self.output_dir / word_filename  # Save to /tmp first
                filename_base = service_name
                output_dir = self.output_dir
                s3_key = None  # Not used in Azure
            else:
                # Local environment: save directly to folder_path (Path object)
                # Ensure folder_path is a valid Path object before using / operator
                if folder_path is None or folder_path == '':
                    # Fallback to output_dir if folder_path is None or empty
                    logger.warning(f"folder_path is None or empty, using output_dir: {self.output_dir}")
                    folder_path = self.output_dir
                elif isinstance(folder_path, str):
                    # Convert string to Path, but check if it's empty
                    if not folder_path.strip():
                        logger.warning(f"folder_path is empty string, using output_dir: {self.output_dir}")
                        folder_path = self.output_dir
                    else:
                        folder_path = Path(folder_path)
                elif not isinstance(folder_path, Path):
                    # If it's still not a Path, fallback to output_dir
                    logger.warning(f"folder_path is not a Path or string, using output_dir: {self.output_dir}")
                    folder_path = self.output_dir
                
                # Final safety check: ensure folder_path is a valid Path
                if not isinstance(folder_path, Path):
                    folder_path = self.output_dir
                
                word_path = folder_path / word_filename
                filename_base = service_name
                output_dir = folder_path
                s3_key = None
        else:
            # Create new document - save to generated_documents (no folder metadata)
            doc_id = str(uuid.uuid4())[:8]
            safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_'))[:50]
            filename_base = f"{safe_title}_{doc_id}"
            word_path = self.output_dir / f"{filename_base}.docx"
            output_dir = self.output_dir
            s3_key = None
        
        # Ensure output directory exists
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # Save Word document
        doc.save(str(word_path))

        # Final safeguard: replace placeholders directly in the saved XML parts
        self._replace_in_saved_docx(str(word_path), {
            'ENTER SERVICE NAME HERE': title,
            'Enter Service Name Here': title,
            'enter service name here': title,
            'Add Title': title,
            '{{SERVICE_NAME}}': title,
        })
        
        # Upload to Azure Blob Storage if in Azure environment
        word_blob_key = None
        pdf_blob_key = None
        if self.use_azure and self.azure_blob_service and (update_metadata or new_proposal_metadata):
            # Construct blob key matching the SharePoint folder structure
            # Format: GCloud {version}/PA Services/Cloud Support Services LOT {lot}/{service_folder}/{filename}
            blob_key = None
            if update_metadata:
                # For Azure, use actual_folder_name which was correctly extracted above
                # For local, folder_path might be a Path, but we should use actual_folder_name for consistency
                folder_name = actual_folder_name
                blob_key = f"GCloud {gcloud_version}/PA Services/Cloud Support Services LOT {update_metadata.get('lot', '3')}/{folder_name}/{word_filename}"
            else:  # new_proposal_metadata
                # Use service_name directly with spaces - NO normalization
                blob_key = f"GCloud {gcloud_version}/PA Services/Cloud Support Services LOT {lot}/{service_name}/{word_filename}"
            
            if blob_key:
                try:
                    self.azure_blob_service.upload_file(word_path, blob_key)
                    word_blob_key = blob_key
                    logger.info(f"Uploaded document to Azure Blob Storage: {word_blob_key}")
                    
                    # Determine PDF blob key
                    pdf_blob_key = blob_key.replace('.docx', '.pdf')
                    
                    # Call Azure PDF converter Function App - ONLY when completing (not saving as draft)
                    # PDFs should only be generated for final completed documents, not drafts
                    pdf_converter_url = os.environ.get("PDF_CONVERTER_FUNCTION_URL")
                    if pdf_converter_url and not save_as_draft:
                        try:
                            import requests  # Import at function level to avoid dependency issues
                            function_key = os.environ.get("PDF_CONVERTER_FUNCTION_KEY", "")
                            
                            # Prepare request payload
                            payload = {
                                "word_blob_key": word_blob_key,
                                "word_container": os.environ.get("AZURE_STORAGE_CONTAINER_NAME", "sharepoint"),
                                "output_container": os.environ.get("AZURE_STORAGE_CONTAINER_NAME", "sharepoint"),
                                "pdf_blob_key": pdf_blob_key
                            }
                            
                            # Build URL with function key if available
                            url = pdf_converter_url
                            if function_key:
                                url = f"{pdf_converter_url}?code={function_key}"
                            
                            # Call PDF converter
                            response = requests.post(
                                url,
                                json=payload,
                                timeout=300  # 5 minute timeout for PDF conversion
                            )
                            
                            if response.status_code == 200:
                                result = response.json()
                                if result.get('success'):
                                    pdf_blob_key = result.get('pdf_blob_key', pdf_blob_key)
                                    logger.info(f"PDF conversion successful: {pdf_blob_key}")
                                else:
                                    logger.warning(f"PDF conversion failed: {result.get('error', 'Unknown error')}")
                                    pdf_blob_key = None
                            else:
                                logger.warning(f"PDF converter returned status {response.status_code}: {response.text}")
                                pdf_blob_key = None
                        except Exception as e:
                            logger.error(f"Failed to call PDF converter: {e}", exc_info=True)
                            pdf_blob_key = None
                    else:
                        logger.warning("PDF_CONVERTER_FUNCTION_URL not set, skipping PDF conversion")
                        pdf_blob_key = None
                except Exception as e:
                    logger.error(f"Failed to upload to Azure Blob Storage: {e}")
                    # Don't fail the entire operation if Azure upload fails
        
        # Upload to S3 if in Lambda environment
        if self.use_s3:
            # Use the S3 key from folder_path if saving to SharePoint, otherwise use generated/
            bucket_sharepoint = os.environ.get('SHAREPOINT_BUCKET_NAME', '')
            bucket_output = os.environ.get('OUTPUT_BUCKET_NAME', '')

            pdf_filename = None
            conversion_bucket = None
            conversion_key = None
            if s3_key:
                word_s3_key = s3_key
                word_bucket = bucket_sharepoint
                pdf_filename = word_filename.replace('.docx', '.pdf')
            else:
                word_s3_key = f"generated/{filename_base}.docx"
                word_bucket = bucket_output
                pdf_filename = f"{filename_base}.pdf"

            if not word_bucket:
                raise ValueError("Target S3 bucket for Word document not configured")

            # Upload document to S3
            import boto3
            s3_client = boto3.client('s3')

            with open(word_path, 'rb') as f:
                s3_client.upload_fileobj(f, word_bucket, word_s3_key)

            # Ensure the PDF converter can access the Word file
            if s3_key and bucket_output:
                # Copy to output bucket for converter access
                conversion_bucket = bucket_output
                conversion_key = f"generated/{filename_base}.docx"
                copy_source = {'Bucket': word_bucket, 'Key': word_s3_key}
                s3_client.copy(copy_source, bucket_output, conversion_key)
            else:
                conversion_bucket = word_bucket
                conversion_key = word_s3_key
            
            # Generate presigned URL for the uploaded document
            word_url = s3_client.generate_presigned_url(
                'get_object',
                Params={'Bucket': word_bucket, 'Key': word_s3_key},
                ExpiresIn=3600  # 1 hour
            )
            
            # Invoke PDF converter Lambda to generate PDF
            pdf_bucket = bucket_output or word_bucket
            pdf_s3_key = f"generated/{pdf_filename}"

            pdf_url = None
            
            try:
                import boto3
                pdf_converter_function = os.environ.get("PDF_CONVERTER_FUNCTION_NAME")
                if pdf_converter_function:
                    lambda_client = boto3.client('lambda')
                    response = lambda_client.invoke(
                        FunctionName=pdf_converter_function,
                        InvocationType='RequestResponse',  # Synchronous invocation
                        Payload=json.dumps({
                            'word_s3_key': conversion_key,
                            'word_bucket': conversion_bucket,
                            'output_bucket': pdf_bucket,
                            'pdf_s3_key': pdf_s3_key
                        })
                    )
                    result = json.loads(response['Payload'].read())
                    if result.get('success'):
                        pdf_url = result.get('pdf_url')
                        pdf_s3_key = result.get('pdf_s3_key', pdf_s3_key)
            except Exception as e:
                # If PDF conversion fails, continue without PDF
                print(f"PDF conversion failed: {e}")
                pdf_url = None
            
            return {
                "word_path": word_url,  # Return presigned URL
                "word_s3_key": word_s3_key,
                "pdf_path": pdf_url or pdf_s3_key,
                "pdf_s3_key": pdf_s3_key,
                "pdf_bucket": pdf_bucket,
                "filename": filename_base
            }
        elif self.use_azure:
            # Azure: return blob keys
            return {
                "word_path": str(word_path),  # Local path for download endpoint
                "word_blob_key": word_blob_key,  # Azure blob key
                "pdf_blob_key": pdf_blob_key,  # Azure PDF blob key (None if conversion failed)
                "pdf_path": pdf_blob_key if pdf_blob_key else "",  # For compatibility
                "filename": filename_base
            }
        else:
            # Docker/local: return local paths
            if update_metadata:
                # For updates, PDF path should be in same folder
                # Ensure output_dir is a valid Path object
                if output_dir is None or output_dir == '':
                    logger.warning(f"output_dir is None or empty for update, using self.output_dir: {self.output_dir}")
                    output_dir = self.output_dir
                elif isinstance(output_dir, str):
                    if not output_dir.strip():
                        logger.warning(f"output_dir is empty string, using self.output_dir: {self.output_dir}")
                        output_dir = self.output_dir
                    else:
                        output_dir = Path(output_dir)
                elif not isinstance(output_dir, Path):
                    logger.warning(f"output_dir is not a Path, using self.output_dir: {self.output_dir}")
                    output_dir = self.output_dir
                
                # Final safety check
                if not isinstance(output_dir, Path):
                    output_dir = self.output_dir
                
                pdf_path = output_dir / f"PA GC{update_metadata.get('gcloud_version', '14')} SERVICE DESC {update_metadata.get('service_name', title)}.pdf"
            else:
                pdf_path = self.output_dir / f"{filename_base}.pdf"
            
            return {
                "word_path": str(word_path),
                "pdf_path": str(pdf_path),
                "filename": filename_base
            }
    
    def _replace_title(self, doc: Document, new_title: str):
        """Replace the first Heading 1 with the new title"""
        replaced = False
        for paragraph in doc.paragraphs:
            if paragraph.style.name == 'Heading 1':
                # Found the title, replace it
                paragraph.text = new_title
                # Preserve formatting
                for run in paragraph.runs:
                    run.font.size = Pt(24)
                    run.font.bold = True
                replaced = True
                break

        # Fallback: replace literal occurrences of common placeholder text in runs
        if not replaced:
            placeholders = ['ENTER SERVICE NAME HERE', 'Enter Service Name Here', 'enter service name here', 'Add Title', '{{SERVICE_NAME}}']
            for p in doc.paragraphs:
                for r in p.runs:
                    for placeholder in placeholders:
                        if placeholder in r.text:
                            r.text = r.text.replace(placeholder, new_title)
                            replaced = True
                            break
                    if replaced:
                        break
                if replaced:
                    break
    
    def _replace_content_sections(
        self,
        doc: Document,
        description: str,
        features: List[str],
        benefits: List[str]
    ):
        """
        Replace content in the template sections
        
        Finds "Short Service Description", "Key Service Features", "Key Service Benefits"
        headings and replaces the content that follows them.
        """
        
        # Track which section we're in
        current_section = None
        section_start_idx = None
        
        found_description = False
        found_features = False
        found_benefits = False

        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            
            # Detect section headings
            if 'Short Service Description' in text:
                current_section = 'description'
                section_start_idx = i
                # Clear existing content under this heading
                self._clear_section_after_heading(doc, i)
                # Replace content after this heading
                self._insert_description(doc, i + 1, description)
                found_description = True
                
            elif 'Key Service Features' in text:
                current_section = 'features'
                section_start_idx = i
                self._clear_section_after_heading(doc, i)
                # Replace content after this heading
                self._insert_bullet_list(doc, i + 1, features)
                found_features = True
                
            elif 'Key Service Benefits' in text:
                current_section = 'benefits'
                section_start_idx = i
                self._clear_section_after_heading(doc, i)
                # Replace content after this heading
                self._insert_bullet_list(doc, i + 1, benefits)
                found_benefits = True

        # Fallback: append sections to end if headings not present
        end_para = doc.paragraphs[-1] if doc.paragraphs else None
        def append_heading(title: str, style: str = 'Heading 2'):
            nonlocal end_para
            end_para = self._insert_paragraph_after(end_para, title, style) if end_para else doc.add_paragraph(title, style)
            return end_para
        if not found_description:
            append_heading('Short Service Description')
            end_para = self._insert_paragraph_after(end_para, description, 'Normal')
        if not found_features:
            append_heading('Key Service Features')
            for f in features:
                end_para = self._insert_paragraph_after(end_para, f, 'List Bullet')
        if not found_benefits:
            append_heading('Key Service Benefits')
            for b in benefits:
                end_para = self._insert_paragraph_after(end_para, b, 'List Bullet')

    def _find_heading_index(self, doc: Document, heading_text: str) -> int | None:
        for i, p in enumerate(doc.paragraphs):
            if p.style and p.style.name.startswith('Heading') and heading_text in (p.text or ''):
                return i
        return None

    def _remove_sections(self, doc: Document, headings: List[str]):
        for h in headings:
            idx = self._find_heading_index(doc, h)
            if idx is not None:
                # Remove from the heading paragraph itself through to next heading
                self._remove_heading_block(doc, idx)

    def _remove_heading_block(self, doc: Document, heading_idx: int):
        body = doc._element.body
        heading_p = doc.paragraphs[heading_idx]._p
        # Remove the heading paragraph itself
        el = heading_p
        # Then continue removing siblings until next heading
        while el is not None:
            nxt = el.getnext()
            if el.tag == qn('w:p'):
                p = Paragraph(el, doc)
                # if we're past the original heading and hit a new heading, stop
                if el is not heading_p and p.style and p.style.name.startswith('Heading'):
                    break
            if el.getparent() is body:
                body.remove(el)
            el = nxt

    def _ensure_toc_and_pagebreak(self, doc: Document):
        # Insert/refresh contents section and get the TOC paragraph
        after_para = self._refresh_contents_section(doc)
        if after_para is not None:
            # Add a page break on the contents paragraph only if not already present
            try:
                has_break = False
                for r in after_para.runs:
                    if r._r.xpath('.//w:br', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                        has_break = True
                        break
                if not has_break:
                    after_para.add_run().add_break(WD_BREAK.PAGE)
            except Exception:
                pass
            return after_para
        # No contents, fallback to last paragraph
        return doc.paragraphs[-1] if doc.paragraphs else None

    def _insert_full_content_block(
        self,
        doc: Document,
        after_para: Paragraph | None,
        service_title: str,
        description: str,
        features: List[str],
        benefits: List[str],
        service_definition: List[dict],
    ) -> Paragraph | None:
        cur = after_para
        def add_para(text: str, style: str | None = None, space_after: int = 12) -> Paragraph:
            nonlocal cur
            cur = self._insert_paragraph_after(cur, text, style) if cur else doc.add_paragraph(text, style)
            try:
                pf = cur.paragraph_format
                pf.space_after = Pt(space_after)
            except Exception:
                pass
            return cur

        # Title as Heading 1 in main content (add extra spacing below)
        add_para(service_title, 'Heading 1', space_after=16)

        # Short Service Description
        add_para('Short Service Description', 'Heading 2', space_after=10)
        add_para(description, 'Normal', space_after=16)

        # Key Service Features as numbered red list (1-10)
        add_para('Key Service Features', 'Heading 2', space_after=10)
        cur = self._insert_numbered_list_block(doc, cur, features[:10])
        # extra spacing after features section
        cur = self._insert_paragraph_after(cur, '')
        try:
            cur.paragraph_format.space_after = Pt(14)
        except Exception:
            pass

        # Key Service Benefits as numbered red list (1-10)
        add_para('Key Service Benefits', 'Heading 2', space_after=8)
        cur = self._insert_numbered_list_block(doc, cur, benefits[:10])
        # extra spacing after benefits section
        cur = self._insert_paragraph_after(cur, '')
        try:
            cur.paragraph_format.space_after = Pt(16)
        except Exception:
            pass

        # Service Definition
        if service_definition:
            add_para('Service Definition', 'Heading 2', space_after=6)
            for block in service_definition:
                subtitle = block.get('subtitle') or ''
                content_html = block.get('content') or ''
                if subtitle:
                    add_para(subtitle, 'Heading 3', space_after=6)
                if content_html:
                    cur = self._insert_html(doc, cur, content_html)
                    # add spacing after html block
                    cur = self._insert_paragraph_after(cur, '')
        return cur

    def _insert_numbered_list_block(self, doc: Document, after_para: Paragraph | None, items: List[str]) -> Paragraph | None:
        """Insert manual numbered list where only the numbers are red and text stays black."""
        cur = after_para
        for idx, item in enumerate(items, start=1):
            # create a normal paragraph and two runs: red number + black text
            cur = self._insert_paragraph_after(cur, '') if cur else doc.add_paragraph('')
            try:
                pf = cur.paragraph_format
                pf.space_after = Pt(4)
            except Exception:
                pass
            r_num = cur.add_run(f"{idx}. ")
            try:
                r_num.font.color.rgb = RGBColor(192, 0, 0)
                r_num.bold = True
            except Exception:
                pass
            r_txt = cur.add_run(item)
            try:
                r_txt.font.color.rgb = RGBColor(0, 0, 0)
            except Exception:
                pass
        return cur

    def _find_heading_element(self, doc: Document, heading_text: str):
        for i, p in enumerate(doc.paragraphs):
            if p.style and p.style.name.startswith('Heading') and heading_text.lower() in (p.text or '').lower():
                return i, p._p
        return None, None

    def _extract_heading_block(self, doc: Document, heading_text: str):
        """Extract a heading and its following content until next heading as XML elements.
        Removes them from the document and returns a list of detached elements to re-append later.
        """
        idx, heading_el = self._find_heading_element(doc, heading_text)
        if heading_el is None:
            return []
        body = doc._element.body
        result = []
        el = heading_el
        while el is not None:
            nxt = el.getnext()
            # Stop before next heading (skip removing it)
            if el is not heading_el and el.tag == qn('w:p'):
                p = Paragraph(el, doc)
                if p.style and p.style.name.startswith('Heading'):
                    break
            # Detach and collect
            result.append(el)
            body.remove(el)
            el = nxt
        return result

    def _extract_block_by_marker(self, doc: Document, marker: str):
        """Find a block that starts where a literal marker appears (in w:t or a:t) and
        extract that element and everything after it to the end of the document.
        Returns a list of detached elements (can be empty if marker not found).
        """
        root = doc._element
        start_el = None
        # Search all text nodes regardless of namespace
        for t in root.xpath(".//*[local-name()='t']"):
            if getattr(t, 'text', None) and marker in t.text:
                # Remove marker text
                t.text = t.text.replace(marker, '')
                # Find top-level ancestor under body
                el = t
                while el is not None and el.tag not in (qn('w:p'), qn('w:tbl'), qn('w:sdt')):
                    el = el.getparent()
                # Climb until direct child of body
                while el is not None and el.getparent() is not None and el.getparent() is not doc._element.body:
                    el = el.getparent()
                start_el = el
                break
        if start_el is None:
            return []
        body = doc._element.body
        result = []
        el = start_el
        while el is not None:
            nxt = el.getnext()
            result.append(el)
            body.remove(el)
            el = nxt
        return result

    def _clear_section_after_heading(self, doc: Document, heading_idx: int):
        """Remove paragraphs following a heading until the next heading or end of document.

        Note: python-docx doesn't support deleting list items as a group, so we remove
        underlying elements paragraph by paragraph.
        """
        # Work at XML level to also remove tables and content controls
        body = doc._element.body
        heading_p = doc.paragraphs[heading_idx]._p
        # Iterate siblings after heading until next heading paragraph
        el = heading_p.getnext()
        while el is not None:
            # Stop at next heading paragraph
            if el.tag == qn('w:p'):
                # Check if it's a heading style
                p = Paragraph(el, doc)
                if p.style and p.style.name.startswith('Heading'):
                    break
            # Remove tables, paragraphs, sdts indiscriminately
            nxt = el.getnext()
            body.remove(el)
            el = nxt
        
    
    def _insert_description(self, doc: Document, start_idx: int, description: str):
        """Insert description text after a heading"""
        # Create a new paragraph right after the heading with the description
        heading_para = doc.paragraphs[start_idx - 1]
        new_para = self._insert_paragraph_after(heading_para, description, 'Normal')
    
    def _insert_bullet_list(self, doc: Document, start_idx: int, items: List[str]):
        """Insert a bullet list after a heading"""
        # Insert bullet list items right after the heading
        heading_para = doc.paragraphs[start_idx - 1]
        insert_after = heading_para
        for item in items:
            insert_after = self._insert_paragraph_after(insert_after, item, 'List Bullet')

    def _insert_service_definition(self, doc: Document, blocks: List[dict]):
        """Insert Service Definition content: subsections with optional images and tables.

        Expected block format examples:
        {"subtitle": "Service Definition Subsection", "content": "Paragraph text...", "images": ["http://..."], "table": [["H1","H2"],["R1C1","R1C2"]] }
        """
        # Find the 'Service Definition' heading
        heading_idx = None
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip() == 'Service Definition' or (
                p.style and p.style.name.startswith('Heading') and 'Service Definition' in p.text
            ):
                heading_idx = i
                break
        if heading_idx is None:
            return
        # Clear existing content under the heading
        self._clear_section_after_heading(doc, heading_idx)

        insert_after = doc.paragraphs[heading_idx]

        # Utilities for images
        def _add_image(after_para, url: str):
            try:
                import requests
                from io import BytesIO
                img_data = requests.get(url, timeout=10).content
                # Insert a new paragraph and add picture to the run
                p = self._insert_paragraph_after(after_para, '')
                run = p.add_run()
                run.add_picture(BytesIO(img_data))
                return p
            except Exception:
                # Ignore image failures silently
                return after_para

        # Build content
        for block in blocks:
            subtitle = block.get('subtitle')
            content_html = block.get('content')
            images = block.get('images', []) or []
            table = block.get('table')

            if subtitle:
                # Sanitize subtitle - replace AI Security advisory with Lorem Ipsum
                if 'AI Security' in subtitle or 'advisory' in subtitle.lower() or '1.1.1' in subtitle or '1.4.1' in subtitle:
                    subtitle = 'Lorem ipsum dolor sit amet'
                insert_after = self._insert_paragraph_after(insert_after, subtitle, 'Heading 3')

            if content_html:
                # Render limited HTML into docx
                insert_after = self._insert_html(doc, insert_after, content_html)

            for img_url in images:
                insert_after = _add_image(insert_after, img_url)

            if table and isinstance(table, list) and table:
                # Insert table after current insert_after by using document-level add and moving near
                rows = len(table)
                cols = max(len(r) for r in table if isinstance(r, list))
                t = doc.add_table(rows=rows, cols=cols)
                t.style = 'Table Grid'
                for r_idx, row in enumerate(table):
                    for c_idx, cell_text in enumerate(row):
                        t.cell(r_idx, c_idx).text = str(cell_text)
                # Add a blank paragraph after table to maintain spacing
                insert_after = self._insert_paragraph_after(insert_after, '')

    def _insert_html(self, doc: Document, after_para, html: str):
        """Very basic HTML renderer supporting <p>, <strong>, <em>, <ul>/<ol>/<li>, <h3>, <br>, <a>.
        Images are expected as separate 'images' array and handled elsewhere.
        Returns the last paragraph inserted for chaining.
        """
        soup = BeautifulSoup(html, 'html.parser')

        def render_inline(run, node):
            if node.name == 'strong' or node.name == 'b':
                r = run.add_text(node.get_text())
                run.bold = True
                return
            if node.name == 'em' or node.name == 'i':
                r = run.add_text(node.get_text())
                run.italic = True
                return
            if node.name == 'br':
                run.add_break()
                return
            # Default text
            run.add_text(node if isinstance(node, str) else node.get_text())

        def add_paragraph_with_inlines(text_or_node, style_name=None):
            p = self._insert_paragraph_after(after_para, '')
            if style_name:
                p.style = style_name
            if isinstance(text_or_node, str):
                p.add_run(text_or_node)
            else:
                for child in text_or_node.children:
                    if isinstance(child, str):
                        p.add_run(child)
                    else:
                        if child.name in ['strong','b','em','i','br']:
                            render_inline(p.add_run(''), child)
                        elif child.name == 'img':
                            src = child.get('src')
                            if src:
                                # Insert image as separate paragraph after current
                                p_img = self._insert_paragraph_after(p, '')
                                try:
                                    from io import BytesIO
                                    if src.startswith('data:'):
                                        # data URL: data:image/png;base64,....
                                        import base64
                                        header, b64 = src.split(',', 1)
                                        img_bytes = base64.b64decode(b64)
                                    else:
                                        import requests
                                        img_bytes = requests.get(src, timeout=10).content
                                    run = p_img.add_run()
                                    run.add_picture(BytesIO(img_bytes))
                                    p = p_img
                                except Exception:
                                    # If image fetch fails, fall back to a link
                                    p_img.add_run(f"[image: {src}]")
                                    p = p_img
                        else:
                            p.add_run(child.get_text())
            return p

        last = after_para
        for el in soup.contents:
            if isinstance(el, str) and el.strip():
                last = add_paragraph_with_inlines(el, None)
            elif getattr(el, 'name', None):
                name = el.name.lower()
                if name == 'h3':
                    last = add_paragraph_with_inlines(el, 'Heading 3')
                elif name == 'p':
                    last = add_paragraph_with_inlines(el, 'Normal')
                elif name in ['ul','ol']:
                    for li in el.find_all('li', recursive=False):
                        p_li = self._insert_paragraph_after(last, li.get_text(), 'List Bullet' if name=='ul' else 'List Number')
                        last = p_li
                elif name == 'br':
                    last = self._insert_paragraph_after(last, '')
                else:
                    # Fallback as paragraph
                    last = add_paragraph_with_inlines(el, 'Normal')
        return last

    def _insert_paragraph_after(self, paragraph: Paragraph, text: str = '', style_name: str | None = None) -> Paragraph:
        """Insert a new paragraph directly after the given paragraph using low-level XML ops."""
        # Create a new paragraph element and insert after current
        new_p = OxmlElement('w:p')
        paragraph._p.addnext(new_p)
        new_para = Paragraph(new_p, paragraph._parent)
        if text:
            new_para.add_run(text)
        if style_name:
            new_para.style = style_name
        return new_para

    def _refresh_contents_section(self, doc: Document):
        """Clear any existing 'Contents' or 'Table of Contents' entries and insert a TOC field.
        Returns the paragraph containing the TOC field if inserted, else None.
        """
        # Find a heading named 'Contents' or 'Table of Contents'
        heading_idx = None
        for i, p in enumerate(doc.paragraphs):
            txt = (p.text or '').strip()
            if txt.lower() in ('contents', 'table of contents') and p.style and p.style.name.startswith('Heading'):
                heading_idx = i
                break
        if heading_idx is None:
            return None
        # Clear everything after this heading until next heading
        self._clear_section_after_heading(doc, heading_idx)
        # Remove any legacy TOC field codes/content and insert a TOC field
        self._remove_existing_toc(doc)
        return self._insert_toc_after_heading(doc, heading_idx)

    def _insert_toc_after_heading(self, doc: Document, heading_idx: int):
        """Insert a Table of Contents field after the specified heading and return the paragraph."""
        para = doc.paragraphs[heading_idx]
        p = self._insert_paragraph_after(para, '')

        # Build field codes: TOC \o "1-3" \h \z \u
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')

        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = 'TOC \\o "1-3" \\h \\z \\u'

        fld_sep = OxmlElement('w:fldChar')
        fld_sep.set(qn('w:fldCharType'), 'separate')

        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')

        r1 = OxmlElement('w:r')
        r1.append(fld_begin)
        r2 = OxmlElement('w:r')
        r2.append(instr)
        r3 = OxmlElement('w:r')
        r3.append(fld_sep)
        r4 = OxmlElement('w:r')
        r4.append(fld_end)

        p._p.append(r1)
        p._p.append(r2)
        p._p.append(r3)
        p._p.append(r4)
        return p

    def _remove_existing_toc(self, doc: Document):
        """Remove any existing TOC field code blocks to avoid duplication."""
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        try:
            for instr in doc._element.xpath('.//w:instrText', namespaces=ns):
                if instr.text and 'TOC' in instr.text:
                    # Remove the containing paragraph or sdt
                    parent = instr
                    # Climb to the paragraph
                    while parent is not None and parent.tag != qn('w:p'):
                        parent = parent.getparent()
                    if parent is not None and parent.getparent() is not None:
                        parent.getparent().remove(parent)
        except Exception:
            pass

    def _enable_update_fields_on_open(self, doc: Document):
        """Enable Word to update fields (e.g., TOC) when the document is opened."""
        try:
            settings = doc.settings
            # Remove existing updateFields if present
            for el in list(settings._element.findall(qn('w:updateFields'))):
                settings._element.remove(el)
            upd = OxmlElement('w:updateFields')
            upd.set(qn('w:val'), 'true')
            settings._element.append(upd)
        except Exception:
            pass
    
    def _update_toc_field(self, doc: Document):
        """Update/refresh the TOC field to populate it with actual headings.
        
        Delete and recreate the TOC field after content insertion.
        This ensures it's populated with the actual headings when LibreOffice/Word processes it.
        """
        try:
            # Find the TOC paragraph
            toc_para = None
            toc_heading_idx = None
            
            for i, p in enumerate(doc.paragraphs):
                txt = (p.text or '').strip()
                if txt.lower() in ('contents', 'table of contents') and p.style and p.style.name.startswith('Heading'):
                    toc_heading_idx = i
                    # Find the paragraph after this heading (should be the TOC field)
                    if i + 1 < len(doc.paragraphs):
                        toc_para = doc.paragraphs[i + 1]
                    break
            
            if toc_para is None or toc_heading_idx is None:
                return
            
            # Delete the existing TOC field paragraph
            try:
                # Remove the TOC paragraph
                toc_para._p.getparent().remove(toc_para._p)
            except Exception:
                pass
            
            # Recreate the TOC field - this will force it to be populated
            # Insert new TOC field after the heading
            heading_para = doc.paragraphs[toc_heading_idx]
            new_toc_para = self._insert_paragraph_after(heading_para, '')
            
            # Build field codes: TOC \o "1-3" \h \z \u
            fld_begin = OxmlElement('w:fldChar')
            fld_begin.set(qn('w:fldCharType'), 'begin')
            
            instr = OxmlElement('w:instrText')
            instr.set(qn('xml:space'), 'preserve')
            instr.text = 'TOC \\o "1-3" \\h \\z \\u'
            
            fld_sep = OxmlElement('w:fldChar')
            fld_sep.set(qn('w:fldCharType'), 'separate')
            
            fld_end = OxmlElement('w:fldChar')
            fld_end.set(qn('w:fldCharType'), 'end')
            
            r1 = OxmlElement('w:r')
            r1.append(fld_begin)
            r2 = OxmlElement('w:r')
            r2.append(instr)
            r3 = OxmlElement('w:r')
            r3.append(fld_sep)
            r4 = OxmlElement('w:r')
            r4.append(fld_end)
            
            new_toc_para._p.append(r1)
            new_toc_para._p.append(r2)
            new_toc_para._p.append(r3)
            new_toc_para._p.append(r4)
            
        except Exception as e:
            # If updating fails, at least ensure updateFields is enabled
            print(f"Warning: Could not recreate TOC field: {e}")
            pass

    def _replace_text_globally(self, doc: Document, old: str, new: str):
        """Aggressively replace text across all document XML parts (covers shapes/textboxes).

        Note: relies on private attributes; safe enough for our constrained use-case.
        """
        try:
            pkg = doc.part.package
            for part in pkg.parts:
                # Only process XML parts
                if hasattr(part, 'blob') and isinstance(part.blob, (bytes, bytearray)):
                    try:
                        xml = part.blob.decode('utf-8')
                    except Exception:
                        continue
                    if old in xml:
                        xml = xml.replace(old, new)
                        part._blob = xml.encode('utf-8')
        except Exception:
            # Fail silently; best-effort replacement
            pass

    def _replace_text_in_all_wt(self, doc: Document, mapping: Dict[str, str]):
        """Replace text in all w:t nodes within the main document (handles shapes/textboxes)."""
        ns = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        }
        try:
            for xpath in ('.//w:t', './/a:t'):
                for t in doc._element.xpath(xpath, namespaces=ns):
                    if t.text:
                        txt = t.text
                        replaced = txt
                        for old, new in mapping.items():
                            if old in replaced:
                                replaced = replaced.replace(old, new)
                        if replaced != txt:
                            t.text = replaced
        except Exception:
            pass

    def _replace_in_saved_docx(self, docx_path: str, mapping: Dict[str, str]):
        """Open the saved .docx and replace placeholders in XML files as a last step."""
        try:
            from zipfile import ZipFile, ZIP_DEFLATED
            import io
            with ZipFile(docx_path, 'r') as zin:
                buf = io.BytesIO()
                with ZipFile(buf, 'w', ZIP_DEFLATED) as zout:
                    for item in zin.infolist():
                        data = zin.read(item.filename)
                        if item.filename.startswith('word/') and item.filename.endswith('.xml'):
                            try:
                                text = data.decode('utf-8')
                                for old, new in mapping.items():
                                    if old in text:
                                        text = text.replace(old, new)
                                data = text.encode('utf-8')
                            except Exception:
                                pass
                        zout.writestr(item, data)
            # Overwrite original file
            with open(docx_path, 'wb') as f:
                f.write(buf.getvalue())
        except Exception:
            pass
    
    def cleanup_old_files(self, days: int = 7):
        """Remove generated documents older than specified days"""
        import time
        current_time = time.time()
        day_seconds = 86400 * days
        
        for file_path in self.output_dir.glob("*"):
            if file_path.is_file():
                file_age = current_time - file_path.stat().st_mtime
                if file_age > day_seconds:
                    file_path.unlink()


# Global instance (lazy initialization - not created in Lambda)
# In Lambda, instances are created in routes with s3_service
# Only create global instance if not in Lambda environment
_use_s3 = os.environ.get("USE_S3", "false").lower() == "true"
if not _use_s3:
    # Only create global instance for Docker/local environment
    document_generator = DocumentGenerator()
else:
    # Lambda environment - don't create global instance
    document_generator = None

